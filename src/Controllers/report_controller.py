import flask
from flask import send_from_directory
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from src.db.entities import Project, db, Host, Vulnerability, Attachment, VulnRef
from itertools import chain
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from time import time
import os
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Report types:
# 1 - per vulnerability
# 2 - per host
# 3 - for current user


UPLOAD_FOLDER = os.path.abspath(os.curdir) + '/upload/'
TMP_FOLDER = os.path.abspath(os.curdir) + '/tmp_reports/'


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(qn(name), value)


def add_page_number(paragraph):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph.add_run()._r.append(fldChar1)
    paragraph.add_run()._r.append(instrText)
    paragraph.add_run()._r.append(fldChar2)


def generate_report(project_id, report_type):
    vulns = Vulnerability.query.filter_by(project_id=project_id).order_by(Vulnerability.id.asc())
    scope = Host.query.filter_by(project_id=project_id).all()

    document = Document()

    add_page_number(document.sections[0].footer.paragraphs[0])

    style = document.styles['Normal']
    font = style.font
    font.size = Pt(14)
    font.name = "Times New Roman"

    document.add_heading('Cyber security evaluation report', 0)
    document.add_page_break()

    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click to update field."
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
    p_element = paragraph._p
    document.add_page_break()

    document.add_heading('Purpose', level=1).style.font.color.rgb = RGBColor(0x0, 0x0, 0x0)
    document.add_paragraph().add_run().add_break()
    document.add_page_break()

    document.add_heading('Summary about system', level=1).style.font.color.rgb = RGBColor(0x0, 0x0, 0x0)
    document.add_paragraph().add_run().add_break()
    p_scope_title = document.add_paragraph()
    p_scope_title.add_run('Scope').bold = True

    scope_table = document.add_table(rows=1, cols=2)
    scope_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    scope_table.style = 'TableGrid'
    hdr_cells = scope_table.rows[0].cells
    hdr_cells[0].text = 'IP'
    hdr_cells[1].text = 'Domain'
    if scope:
        for host in scope:
            row_cells = scope_table.add_row().cells
            row_cells[0].text = host.ip or "-"
            row_cells[1].text = host.domain or "-"
    document.add_page_break()

    document.add_heading('Results', level=1).style.font.color.rgb = RGBColor(0x0, 0x0, 0x0)

    if report_type == 1:
        for vuln in vulns:
            refs = db.session.query(VulnRef.host_id).filter(VulnRef.vuln_id == vuln.id).all()
            hosts = Host.query.filter(Host.project_id == project_id,
                                      Host.id.in_(list(chain.from_iterable(refs)))).all()
            attachments = Attachment.query.filter(Attachment.vuln_id == vuln.id)

            vuln_title = document.add_heading(vuln.name, level=2)
            vuln_title.style.font.size = Pt(14)
            vuln_title.style.font.color.rgb = RGBColor(0, 0, 0)
            vuln_title.add_run().add_break()

            host_num = 1
            if len(hosts) > 0:
                target = document.add_paragraph()
                target.add_run('Target: ').bold = True
                for host in hosts:
                    target.add_run(host.ip or host.domain)
                    if host_num != len(hosts):
                        target.add_run(", ")
                        host_num += 1

            p_path = document.add_paragraph()
            p_path.add_run('Path: ').bold = True
            p_path.add_run(vuln.full_path)
            p_path.add_run().add_break()

            vuln_settings = {
                0: {
                    "text": "info",
                    "R": 255,
                    "G": 0,
                    "B": 0,
                },
                1: {
                    "text": "low",
                    "R": 0,
                    "G": 176,
                    "B": 80,
                },
                2: {
                    "text": "medium",
                    "R": 255,
                    "G": 192,
                    "B": 0,
                },
                3: {
                    "text": "high",
                    "R": 255,
                    "G": 0,
                    "B": 0,
                },
            }

            criticality = document.add_paragraph()
            criticality.add_run('Criticality: ').bold = True
            criticality.add_run(vuln_settings[vuln.criticality]["text"]).font.color.rgb = RGBColor(
                int(vuln_settings[vuln.criticality]["R"]),
                int(vuln_settings[vuln.criticality]["G"]),
                int(vuln_settings[vuln.criticality]["B"])
            )
            probability = document.add_paragraph()
            probability.add_run('Probability: ').bold = True
            probability.add_run(vuln_settings[vuln.probability]["text"]).font.color.rgb = RGBColor(
                int(vuln_settings[vuln.probability]["R"]),
                int(vuln_settings[vuln.probability]["G"]),
                int(vuln_settings[vuln.probability]["B"])
            )
            final_criticality = document.add_paragraph()
            final_criticality.add_run('Final criticality: ').bold = True
            final_criticality.add_run(vuln_settings[vuln.final_criticality]["text"]).font.color.rgb = RGBColor(
                int(vuln_settings[vuln.final_criticality]["R"]),
                int(vuln_settings[vuln.final_criticality]["G"]),
                int(vuln_settings[vuln.final_criticality]["B"])
            )
            document.add_paragraph().add_run().add_break()

            p_description_title = document.add_paragraph()
            p_description_title.add_run('Description').bold = True
            document.add_paragraph(vuln.description).add_run().add_break()

            p_risk_title = document.add_paragraph()
            p_risk_title.add_run('Risk').bold = True
            document.add_paragraph(vuln.risk).add_run().add_break()

            p_details_title = document.add_paragraph()
            p_details_title.add_run('Technical details').bold = True
            document.add_paragraph(vuln.details).add_run().add_break()

            p_recommendation_title = document.add_paragraph()
            p_recommendation_title.add_run('Recommendation').bold = True
            document.add_paragraph(vuln.recommendation).add_run().add_break()

            if attachments:
                for attach in attachments:
                    document.add_picture(os.path.join(UPLOAD_FOLDER + str(vuln.id), attach.filename), width=Inches(6))
                    picture_paragraph = document.paragraphs[-1]
                    picture_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    document.add_paragraph("Fig.  - " + attach.description).add_run()
                    pic_desc_paragraph = document.paragraphs[-1]
                    pic_desc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    document.add_paragraph().add_run().add_break()

        tmp_filename = 'report' + str(int(time())) + '.docx'
        document.save(os.path.join(TMP_FOLDER, tmp_filename))
        return flask.send_from_directory(TMP_FOLDER, tmp_filename, as_attachment=True, cache_timeout=0)

    elif report_type == 2:
        for vuln in vulns:
            refs = db.session.query(VulnRef.host_id).filter(VulnRef.vuln_id == vuln.id).all()
            hosts = Host.query.filter(Host.project_id == project_id,
                                      Host.id.in_(list(chain.from_iterable(refs)))).all()
            attachments = Attachment.query.filter(Attachment.vuln_id == vuln.id)

            if hosts and len(hosts) > 0:
                for host in hosts:
                    vuln_title = document.add_heading(vuln.name + " - ", level=2)
                    vuln_title.add_run(host.ip or host.domain)
                    vuln_title.style.font.size = Pt(14)
                    vuln_title.style.font.color.rgb = RGBColor(0, 0, 0)
                    vuln_title.add_run().add_break()

                    target = document.add_paragraph()
                    target.add_run('Target: ').bold = True
                    target.add_run(host.ip or host.domain)

                    p_path = document.add_paragraph()
                    p_path.add_run('Path: ').bold = True
                    p_path.add_run(vuln.full_path)
                    p_path.add_run().add_break()

                    vuln_settings = {
                        0: {
                            "text": "info",
                            "R": 255,
                            "G": 0,
                            "B": 0,
                        },
                        1: {
                            "text": "low",
                            "R": 0,
                            "G": 176,
                            "B": 80,
                        },
                        2: {
                            "text": "medium",
                            "R": 255,
                            "G": 192,
                            "B": 0,
                        },
                        3: {
                            "text": "high",
                            "R": 255,
                            "G": 0,
                            "B": 0,
                        },
                    }

                    criticality = document.add_paragraph()
                    criticality.add_run('Criticality: ').bold = True
                    criticality.add_run(vuln_settings[vuln.criticality]["text"]).font.color.rgb = RGBColor(
                        int(vuln_settings[vuln.criticality]["R"]),
                        int(vuln_settings[vuln.criticality]["G"]),
                        int(vuln_settings[vuln.criticality]["B"])
                    )
                    probability = document.add_paragraph()
                    probability.add_run('Probability: ').bold = True
                    probability.add_run(vuln_settings[vuln.probability]["text"]).font.color.rgb = RGBColor(
                        int(vuln_settings[vuln.probability]["R"]),
                        int(vuln_settings[vuln.probability]["G"]),
                        int(vuln_settings[vuln.probability]["B"])
                    )
                    final_criticality = document.add_paragraph()
                    final_criticality.add_run('Final criticality: ').bold = True
                    final_criticality.add_run(vuln_settings[vuln.final_criticality]["text"]).font.color.rgb = RGBColor(
                        int(vuln_settings[vuln.final_criticality]["R"]),
                        int(vuln_settings[vuln.final_criticality]["G"]),
                        int(vuln_settings[vuln.final_criticality]["B"])
                    )
                    document.add_paragraph().add_run().add_break()

                    p_description_title = document.add_paragraph()
                    p_description_title.add_run('Description').bold = True
                    document.add_paragraph(vuln.description).add_run().add_break()

                    p_risk_title = document.add_paragraph()
                    p_risk_title.add_run('Risk').bold = True
                    document.add_paragraph(vuln.risk).add_run().add_break()

                    p_details_title = document.add_paragraph()
                    p_details_title.add_run('Technical details').bold = True
                    document.add_paragraph(vuln.details).add_run().add_break()

                    p_recommendation_title = document.add_paragraph()
                    p_recommendation_title.add_run('Recommendation').bold = True
                    document.add_paragraph(vuln.recommendation).add_run().add_break()

                    if attachments:
                        for attach in attachments:
                            document.add_picture(os.path.join(UPLOAD_FOLDER + str(vuln.id), attach.filename),
                                                 width=Inches(6))
                            picture_paragraph = document.paragraphs[-1]
                            picture_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            document.add_paragraph("Fig.  - " + attach.description).add_run()
                            pic_desc_paragraph = document.paragraphs[-1]
                            pic_desc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            document.add_paragraph().add_run().add_break()
            else:
                vuln_title = document.add_heading(vuln.name, level=2)
                vuln_title.style.font.size = Pt(14)
                vuln_title.style.font.color.rgb = RGBColor(0, 0, 0)
                vuln_title.add_run().add_break()

                p_path = document.add_paragraph()
                p_path.add_run('Path: ').bold = True
                p_path.add_run(vuln.full_path)
                p_path.add_run().add_break()

                vuln_settings = {
                    0: {
                        "text": "info",
                        "R": 255,
                        "G": 0,
                        "B": 0,
                    },
                    1: {
                        "text": "low",
                        "R": 0,
                        "G": 176,
                        "B": 80,
                    },
                    2: {
                        "text": "medium",
                        "R": 255,
                        "G": 192,
                        "B": 0,
                    },
                    3: {
                        "text": "high",
                        "R": 255,
                        "G": 0,
                        "B": 0,
                    },
                }

                criticality = document.add_paragraph()
                criticality.add_run('Criticality: ').bold = True
                criticality.add_run(vuln_settings[vuln.criticality]["text"]).font.color.rgb = RGBColor(
                    int(vuln_settings[vuln.criticality]["R"]),
                    int(vuln_settings[vuln.criticality]["G"]),
                    int(vuln_settings[vuln.criticality]["B"])
                )
                probability = document.add_paragraph()
                probability.add_run('Probability: ').bold = True
                probability.add_run(vuln_settings[vuln.probability]["text"]).font.color.rgb = RGBColor(
                    int(vuln_settings[vuln.probability]["R"]),
                    int(vuln_settings[vuln.probability]["G"]),
                    int(vuln_settings[vuln.probability]["B"])
                )
                final_criticality = document.add_paragraph()
                final_criticality.add_run('Final criticality: ').bold = True
                final_criticality.add_run(vuln_settings[vuln.final_criticality]["text"]).font.color.rgb = RGBColor(
                    int(vuln_settings[vuln.final_criticality]["R"]),
                    int(vuln_settings[vuln.final_criticality]["G"]),
                    int(vuln_settings[vuln.final_criticality]["B"])
                )
                document.add_paragraph().add_run().add_break()

                p_description_title = document.add_paragraph()
                p_description_title.add_run('Description').bold = True
                document.add_paragraph(vuln.description).add_run().add_break()

                p_risk_title = document.add_paragraph()
                p_risk_title.add_run('Risk').bold = True
                document.add_paragraph(vuln.risk).add_run().add_break()

                p_details_title = document.add_paragraph()
                p_details_title.add_run('Technical details').bold = True
                document.add_paragraph(vuln.details).add_run().add_break()

                p_recommendation_title = document.add_paragraph()
                p_recommendation_title.add_run('Recommendation').bold = True
                document.add_paragraph(vuln.recommendation).add_run().add_break()

                if attachments:
                    for attach in attachments:
                        document.add_picture(os.path.join(UPLOAD_FOLDER + str(vuln.id), attach.filename),
                                             width=Inches(6))
                        picture_paragraph = document.paragraphs[-1]
                        picture_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        document.add_paragraph("Fig.  - " + attach.description).add_run()
                        pic_desc_paragraph = document.paragraphs[-1]
                        pic_desc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        document.add_paragraph().add_run().add_break()

        tmp_filename = 'report' + str(int(time())) + '.docx'
        document.save(os.path.join(TMP_FOLDER, tmp_filename))
        return flask.send_from_directory(TMP_FOLDER, tmp_filename, as_attachment=True, cache_timeout=0)
